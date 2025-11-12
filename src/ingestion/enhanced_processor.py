"""
Enhanced document processing module with OCR support.
Handles PDFs (text and scanned), DOCX, and extracts geographic metadata.
"""

import os
import re
from pathlib import Path
from typing import List, Dict, Optional
from dataclasses import dataclass
import logging

import pypdf
from tqdm import tqdm

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Zimbabwe geographic entities
ZIMBABWE_PROVINCES = [
    "Harare", "Bulawayo", "Mashonaland Central", "Mashonaland East", 
    "Mashonaland West", "Manicaland", "Masvingo", "Midlands",
    "Matabeleland North", "Matabeleland South"
]

ZIMBABWE_DISTRICTS = [
    "Bindura", "Hwange", "Mutare", "Gweru", "Masvingo", "Goromonzi",
    "Chipinge", "Beitbridge", "Gokwe", "Tsholotsho", "Binga", "Bikita",
    "Bubi", "Buhera", "Bulilima", "Chegutu", "Chikomba", "Chimanimani",
    "Chiredzi", "Chirumhanzu", "Chivi", "Guruve", "Gutu", "Gwanda",
    "Hwedza", "Insiza", "Kariba", "Kwekwe", "Lupane", "Makonde",
    "Mangwe", "Marondera", "Matobo", "Mazowe", "Mberengwa", "Mbire",
    "Mhondoro-Ngezi", "Mudzi", "Murehwa", "Mutasa", "Mutoko", "Mwenezi",
    "Nyanga", "Rushinga", "Sanyati", "Seke", "Shamva", "Shurugwi",
    "UMP", "Umguza", "Umzingwane", "Zaka", "Zvimba", "Zvishavane",
    "Darwin", "Centenary"
]

NATURAL_REGIONS = ["Region I", "Region II", "Region III", "Region IV", "Region V",
                   "Region IIa", "Region IIb"]


@dataclass
class Document:
    """Represents a processed document chunk."""
    content: str
    metadata: Dict
    doc_id: str
    chunk_id: int


class EnhancedDocumentProcessor:
    """
    Enhanced processor for agricultural documents with OCR and geographic extraction.
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200, use_ocr: bool = True):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.use_ocr = use_ocr
        
        # Check OCR availability
        self.ocr_available = False
        if use_ocr:
            try:
                import pytesseract
                from pdf2image import convert_from_path
                self.ocr_available = True
                logger.info("OCR support enabled")
            except ImportError:
                logger.warning("OCR libraries not available. Install: pip install pytesseract pdf2image pillow")
                logger.warning("Also install Tesseract: brew install tesseract (Mac) or apt install tesseract-ocr (Linux)")
        
        # Check DOCX support
        self.docx_available = False
        try:
            import docx
            self.docx_available = True
            logger.info("DOCX support enabled")
        except ImportError:
            logger.warning("DOCX support not available. Install: pip install python-docx")
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF file with fallback to OCR."""
        text = ""
        
        try:
            # First try regular text extraction
            with open(pdf_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # If extracted text is too short, it might be scanned - try OCR
            if len(text.strip()) < 100 and self.ocr_available:
                logger.info(f"Text extraction yielded little content, attempting OCR for {Path(pdf_path).name}")
                text = self._extract_with_ocr(pdf_path)
                
        except Exception as e:
            logger.error(f"Error extracting text from {pdf_path}: {e}")
            # Try OCR as fallback
            if self.ocr_available:
                try:
                    text = self._extract_with_ocr(pdf_path)
                except Exception as ocr_error:
                    logger.error(f"OCR also failed: {ocr_error}")
        
        return text
    
    def _extract_with_ocr(self, pdf_path: str) -> str:
        """Extract text using OCR (for scanned PDFs)."""
        if not self.ocr_available:
            return ""
        
        try:
            import pytesseract
            from pdf2image import convert_from_path
            from PIL import Image
            
            logger.info(f"Running OCR on {Path(pdf_path).name}...")
            
            # Convert PDF to images
            images = convert_from_path(pdf_path, dpi=300, first_page=1, last_page=10)  # Limit to first 10 pages
            
            text = ""
            for i, image in enumerate(images):
                # Extract text from image
                page_text = pytesseract.image_to_string(image)
                text += f"\n--- Page {i+1} ---\n{page_text}\n"
            
            logger.info(f"OCR completed for {Path(pdf_path).name}, extracted {len(text)} characters")
            return text
            
        except Exception as e:
            logger.error(f"OCR failed for {pdf_path}: {e}")
            return ""
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file."""
        if not self.docx_available:
            logger.warning(f"Cannot process DOCX file {docx_path} - python-docx not installed")
            return ""
        
        try:
            import docx
            doc = docx.Document(docx_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {docx_path}: {e}")
            return ""
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\.\,\-\:\;\?\!\\/\(\)\%\$]', '', text)
        
        # Remove page numbers and headers/footers patterns
        text = re.sub(r'\b\d{1,3}\s*\|\s*Page\b', '', text)
        text = re.sub(r'\bPage\s+\d{1,3}\b', '', text)
        
        return text.strip()
    
    def extract_geographic_metadata(self, text: str, filename: str) -> Dict:
        """Extract geographic information (provinces, districts, regions) from text and filename."""
        text_lower = text.lower()
        filename_lower = filename.lower()
        
        metadata = {
            'provinces': [],
            'districts': [],
            'natural_regions': [],
            'cities': []
        }
        
        # Check provinces
        for province in ZIMBABWE_PROVINCES:
            if province.lower() in text_lower or province.lower() in filename_lower:
                metadata['provinces'].append(province)
        
        # Check districts
        for district in ZIMBABWE_DISTRICTS:
            if district.lower() in text_lower or district.lower() in filename_lower:
                metadata['districts'].append(district)
        
        # Check natural regions
        for region in NATURAL_REGIONS:
            if region.lower() in text_lower:
                metadata['natural_regions'].append(region)
        
        # Remove duplicates
        metadata['provinces'] = list(set(metadata['provinces']))
        metadata['districts'] = list(set(metadata['districts']))
        metadata['natural_regions'] = list(set(metadata['natural_regions']))
        
        return metadata
    
    def chunk_text(self, text: str, doc_id: str) -> List[Document]:
        """Split text into overlapping chunks."""
        chunks = []
        words = text.split()
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            if len(chunk_text) < 100:  # Skip very small chunks
                continue
                
            chunks.append(Document(
                content=chunk_text,
                metadata={
                    'source': doc_id,
                    'chunk_index': len(chunks),
                    'total_words': len(chunk_words)
                },
                doc_id=doc_id,
                chunk_id=len(chunks)
            ))
        
        return chunks
    
    def extract_metadata(self, file_path: str, text: str) -> Dict:
        """Extract comprehensive metadata from document."""
        file_name = Path(file_path).stem
        
        # Determine document category based on filename
        categories = {
            'crop': ['maize', 'soya', 'bean', 'pepper', 'grains', 'horticultur', 'wheat', 'potato'],
            'livestock': ['cattle', 'goat', 'pig', 'broiler', 'chicken', 'fish', 'dairy', 'beef'],
            'policy': ['policy', 'framework', 'strategy', 'vision', 'zimasset'],
            'climate': ['climate', 'csa', 'ipcc', 'weather'],
            'market': ['market', 'value-chain', 'linkage', 'trade'],
            'food_security': ['zimvac', 'fews', 'livelihoods', 'assessment', 'nutrition'],
            'district_profile': ['district-profile', 'district profile']
        }
        
        category = 'general'
        for cat, keywords in categories.items():
            if any(keyword in file_name.lower() for keyword in keywords):
                category = cat
                break
        
        # Extract geographic metadata
        geo_metadata = self.extract_geographic_metadata(text, file_name)
        
        metadata = {
            'filename': Path(file_path).name,
            'category': category,
            'file_path': file_path,
            **geo_metadata
        }
        
        return metadata
    
    def process_document(self, file_path: str) -> List[Document]:
        """Process a single document (PDF or DOCX)."""
        logger.info(f"Processing: {Path(file_path).name}")
        
        file_ext = Path(file_path).suffix.lower()
        
        # Extract text based on file type
        if file_ext == '.pdf':
            raw_text = self.extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            raw_text = self.extract_text_from_docx(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_ext} for {file_path}")
            return []
        
        if not raw_text or len(raw_text.strip()) < 50:
            logger.warning(f"No meaningful text extracted from {file_path}")
            return []
        
        # Clean text
        clean_text = self.clean_text(raw_text)
        
        # Extract metadata (including geographic info)
        metadata = self.extract_metadata(file_path, clean_text)
        doc_id = Path(file_path).stem
        
        # Chunk text
        chunks = self.chunk_text(clean_text, doc_id)
        
        # Add metadata to all chunks
        for chunk in chunks:
            chunk.metadata.update(metadata)
        
        logger.info(f"Created {len(chunks)} chunks from {Path(file_path).name}")
        if metadata['districts']:
            logger.info(f"  Districts found: {', '.join(metadata['districts'])}")
        if metadata['provinces']:
            logger.info(f"  Provinces found: {', '.join(metadata['provinces'])}")
        
        return chunks
    
    def process_directory(self, directory_path: str) -> List[Document]:
        """Process all PDF and DOCX documents in a directory."""
        all_chunks = []
        
        # Find all supported files
        supported_extensions = ['*.pdf', '*.docx', '*.doc']
        all_files = []
        for ext in supported_extensions:
            all_files.extend(list(Path(directory_path).glob(ext)))
        
        logger.info(f"Found {len(all_files)} files to process")
        
        for file_path in tqdm(all_files, desc="Processing documents"):
            try:
                chunks = self.process_document(str(file_path))
                all_chunks.extend(chunks)
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {e}")
                continue
        
        logger.info(f"Total chunks created: {len(all_chunks)}")
        return all_chunks


if __name__ == "__main__":
    # Test the processor
    processor = EnhancedDocumentProcessor()
    data_path = "/Users/providencemtendereki/Documents/Old Staff/PDFs data"
    
    if os.path.exists(data_path):
        # Test on first few files
        test_files = list(Path(data_path).glob("*-District-Profile.pdf"))[:3]
        for test_file in test_files:
            print(f"\nTesting: {test_file.name}")
            chunks = processor.process_document(str(test_file))
            print(f"Created {len(chunks)} chunks")
            if chunks:
                print(f"Metadata: {chunks[0].metadata}")
    else:
        print(f"Data path not found: {data_path}")
